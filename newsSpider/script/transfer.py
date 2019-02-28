import js2py
import requests
import re
import json
from docx import Document
from docx.shared import Inches

# 网页版有个跟内容加密sign参数，需要计算。为了方便，所以选用手机版的网址进行翻译
# url = "https://fanyi.baidu.com/v2transapi"
queryList = ["KIEV, Ukraine (AP) — The Ukrainian president has urged NATO to deploy naval ships to the Sea of Azov amid a standoff with Russia, a statement sharply criticized by the Kremlin as a provocation aimed at further inflaming tensions.",
"President Petro Poroshenko made the call in an interview with the German daily Bild published Thursday, saying that “Germany is one of our closest allies and we hope that states within NATO are now ready to relocate naval ships to the Sea of Azov in order to assist Ukraine and provide security.”"
]
imageList = ["timqwqwqwvg.jpg","timqwqwqwvg.jpg"]

headers = {
    'Accept':'*/*',
    'Content-Type':'application/x-www-form-urlencoded',
    'Cookie':'BAIDUCUID=++; PSTM=1491655249; BIDUPSID=1ECA43D665529765134F5CCD000FB51A; REALTIME_TRANS_SWITCH=1; FANYI_WORD_SWITCH=1; HISTORY_SWITCH=1; SOUND_SPD_SWITCH=1; SOUND_PREFER_SWITCH=1; hasSeenTips=1; BAIDUID=A57974B94365D8C0357D2613B4050312:FG=1; to_lang_often=%5B%7B%22value%22%3A%22en%22%2C%22text%22%3A%22%u82F1%u8BED%22%7D%2C%7B%22value%22%3A%22zh%22%2C%22text%22%3A%22%u4E2D%u6587%22%7D%5D; from_lang_often=%5B%7B%22value%22%3A%22zh%22%2C%22text%22%3A%22%u4E2D%u6587%22%7D%2C%7B%22value%22%3A%22en%22%2C%22text%22%3A%22%u82F1%u8BED%22%7D%5D; MCITY=-158%3A; BDUSS=kZLaWJsVDdCLVk4YklFR3ZJbGt-OFdrZXFLTGxBSkNZQXRzTGZTRFlsaThGaHBjQVFBQUFBJCQAAAAAAAAAAAEAAABQ6oU00NDV37OjyN0AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAALyJ8lu8ifJbW; BDORZ=B490B5EBF6F3CD402E515D22BCDA1598; H_PS_PSSID=1422_21111_26350_27889_22074; BDSFRCVID=nbLOJeC626l72rQ7pkIIuFMPK2zRHN3TH6aoHPR1IWdIjKlHiLPxEG0PjU8g0Kubh02GogKKLmOTHpKF_2uxOjjg8UtVJeC6EG0P3J; H_BDCLCKID_SF=tJPHoDI-JKL3j5ruM-rV-JD0-fTBa4oXHD7yWCv8KqRcOR5Jj6K-h5bXKRLfbPnAW6Tl2xJcttcCOhuR3MA--t4n0bKDBlbbWCQiWpnVBUJWsq0x0-6le-bQypoaaT8H3KOMahv1al7xO-JoQlPK5JkgMx6MqpQJQeQ-5KQN3KJmhpFu-n5jHjj0Da_83H; delPer=0; PSINO=7; locale=zh; BDRCVFR[feWj1Vr5u3D]=I67x6TjHwwYf0; Hm_lvt_afd111fa62852d1f37001d1f980b6800=1543476095; Hm_lpvt_afd111fa62852d1f37001d1f980b6800=1543476095; Hm_lvt_64ecd82404c51e03dc91cb9e8c025574=1541765799,1543472813,1543476096; Hm_lpvt_64ecd82404c51e03dc91cb9e8c025574=1543476096',
    'Host':'fanyi.baidu.com',
    'User-Agent':'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_13_4) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/66.0.3359.139 Safari/537.36'
    }

document = Document()
document.add_heading('Document Title', level=1)
i = 0
for query in queryList:
    i += 1
    # 计算sign
    session = requests.Session()
    session.headers = headers
    response = session.get("https://fanyi.baidu.com")
    gtk = re.findall(";window.gtk = ('.*?');",response.content.decode())[0]

    context = js2py.EvalJs()
    js = r'''
    function a(r) {
            if (Array.isArray(r)) {
                for (var o = 0, t = Array(r.length); o < r.length; o++)t[o] = r[o];
                return t
            } return Array.from(r)
        }
    function n(r, o) {
            for (var t = 0; t < o.length - 2; t += 3) {
                var a = o.charAt(t + 2); a = a >= "a" ? a.charCodeAt(0) - 87 : Number(a), a = "+" === o.charAt(t + 1) ? r >>> a : r << a, r = "+" === o.charAt(t) ? r + a & 4294967295 : r ^ a
            } return r
        } 
    function e(r) {
            var o = r.match(/[\uD800-\uDBFF][\uDC00-\uDFFF]/g); if (null === o) {
                var t = r.length; t > 30 && (r = "" + r.substr(0, 10) + r.substr(Math.floor(t / 2) - 5, 10) + r.substr(-10, 10))
            } else {
                for (var e = r.split(/[\uD800-\uDBFF][\uDC00-\uDFFF]/), C = 0, h = e.length, f = []; h > C; C++)"" !== e[C] && f.push.apply(f, a(e[C].split(""))), C !== h - 1 && f.push(o[C]);
                var g = f.length; g > 30 && (r = f.slice(0, 10).join("") + f.slice(Math.floor(g / 2) - 5, Math.floor(g / 2) + 5).join("") + f.slice(-10).join(""))
            } var u = void 0, l = "" + String.fromCharCode(103) + String.fromCharCode(116) + String.fromCharCode(107); u = null !== i ? i : (i = window[l] || "") || "";
            for (var d = u.split("."), m = Number(d[0]) || 0, s = Number(d[1]) || 0, S = [], c = 0, v = 0; v < r.length; v++) {
                var A = r.charCodeAt(v); 128 > A ? S[c++] = A : (2048 > A ? S[c++] = A >> 6 | 192 : (55296 === (64512 & A) && v + 1 < r.length && 56320 === (64512 & r.charCodeAt(v + 1)) ? (A = 65536 + ((1023 & A) << 10) + (1023 & r.charCodeAt(++v)), S[c++] = A >> 18 | 240, S[c++] = A >> 12 & 63 | 128) : S[c++] = A >> 12 | 224, S[c++] = A >> 6 & 63 | 128), S[c++] = 63 & A | 128)
            } for (var p = m, F = "" + String.fromCharCode(43) + String.fromCharCode(45) + String.fromCharCode(97) + ("" + String.fromCharCode(94) + String.fromCharCode(43) + String.fromCharCode(54)), D = "" + String.fromCharCode(43) + String.fromCharCode(45) + String.fromCharCode(51) + ("" + String.fromCharCode(94) + String.fromCharCode(43) + String.fromCharCode(98)) + ("" + String.fromCharCode(43) + String.fromCharCode(45) + String.fromCharCode(102)), b = 0; b < S.length; b++)p += S[b], p = n(p, F);
            return p = n(p, D), p ^= s, 0 > p && (p = (2147483647 & p) + 2147483648), p %= 1e6, p.toString() + "." + (p ^ m)
        } 
        var i = "320305.131321201"

    '''
    #js中添加一行gtk
    js = js.replace('\'null !== i ? i : (i = window[l] || "") || ""\'',gtk)
    # print(js)
    #执行js
    context.execute(js)
    #调用函数得到sign
    sign = context.e(query)
    v2transapi = 'https://fanyi.baidu.com/v2transapi?from=%s&to=%s&query=%s' \
             '&transtype=translang&simple_means_flag=3&sign=%s&token=%s' % ('en', 'zh', query, sign, '4c8a3fcf0884d0f90b82d9325c6c7357')
    print(v2transapi)
    cook = 'BAIDUID=7F3F75BED2B79B21463A426DFE2C90E6:FG=1; BIDUPSID=7F3F75BED2B79B21463A426DFE2C90E6; PSTM=1550388280; BDUSS=djUFFPUDkwSXg2SDRYS1lsMjBSZGRWVnp5SmhEWGdqVk8wYXZQQmFnbEptWkJjQVFBQUFBJCQAAAAAAAAAAAEAAABQ6oU00NDV37OjyN0AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAEkMaVxJDGlcRm; REALTIME_TRANS_SWITCH=1; FANYI_WORD_SWITCH=1; HISTORY_SWITCH=1; SOUND_SPD_SWITCH=1; SOUND_PREFER_SWITCH=1; from_lang_often=%5B%7B%22value%22%3A%22zh%22%2C%22text%22%3A%22%u4E2D%u6587%22%7D%2C%7B%22value%22%3A%22en%22%2C%22text%22%3A%22%u82F1%u8BED%22%7D%5D; to_lang_often=%5B%7B%22value%22%3A%22en%22%2C%22text%22%3A%22%u82F1%u8BED%22%7D%2C%7B%22value%22%3A%22zh%22%2C%22text%22%3A%22%u4E2D%u6587%22%7D%5D; BDORZ=B490B5EBF6F3CD402E515D22BCDA1598; H_PS_PSSID=1421_25809_21116_26350_28413_22157; delPer=0; locale=zh; Hm_lvt_afd111fa62852d1f37001d1f980b6800=1551338887; OUTFOX_SEARCH_USER_ID_NCOO=1668143724.1957805; Hm_lpvt_afd111fa62852d1f37001d1f980b6800=1551339452; ___rl__test__cookies=1551339454571; Hm_lvt_64ecd82404c51e03dc91cb9e8c025574=1550660009,1550797522,1551338876,1551339476; PSINO=7; Hm_lpvt_64ecd82404c51e03dc91cb9e8c025574=1551340541; BDRCVFR[feWj1Vr5u3D]=I67x6TjHwwYf0'
    #申明一个用于存储手动cookies的字典
    manual_cookies={}
    cookies_txt = cook.strip(';')
    #手动分割添加cookie
    for item in cookies_txt.split(';'):
        name,value=item.strip().split('=',1)
        manual_cookies[name]=value

    #将字典转为CookieJar：
    cookiesJar = requests.utils.cookiejar_from_dict(manual_cookies, cookiejar=None,overwrite=True)
    r = requests.get(v2transapi,cookies=cookiesJar).json()
    # param = {
    #     'from':'en',
    #     'to':'zh',
    #     'query':query,
    #     'sign':sign,
    #     'token':'6fcc049803b1d94d08e03b38ca7b251d',
    #     'transtype': 'realtime',
    #     'simple_means_flag': '3'
    # }
    # r = requests.post(url=url,data=param,headers=headers).json()
    print(r)
    result = r['trans_result']['data'][0]['dst']
    p = document.add_paragraph(result)
    if i % 3 == 0 and len(imageList) > 0:
        document.add_picture(imageList.pop(), width=Inches(1.25))

document.save('demo.docx')
