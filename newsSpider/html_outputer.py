#!/usr/bin/env python2
# -*- coding: UTF-8 -*-
# import pymysql.cursors
import requests
import json
from docx import Document
from docx.shared import Inches
import time
import os
from newsSpider import html_downloader

class HtmlOutputer(object):

    def __init__(self):
        self.datas = []
        self.downloader = html_downloader.HtmlDownloader()

    def collect_data(self, data):
        if data is None:
            return
        self.datas.append(data)

    def output_word(self, data):
        if data is None:
            return
        # print(data)
        
        queryList = data['text']
        imageList = data['images_src']
        title = data['title']

        document = Document()
        param = {
            'from':'en',
            'to':'zh',
            'query':title,
        }
        title_ch = self.downloader.baidu_transfer(param)
        document.add_heading(title_ch, level=1)
        i = 0
        for query in queryList:
            # time.sleep(1)
            i += 1
            param = {
                'from':'en',
                'to':'zh',
                'query':query,
            }
            result = self.downloader.baidu_transfer(param)
            p = document.add_paragraph(result)
            print(result)
            if i % 3 == 0 and len(imageList) > 0:
                document.add_picture(imageList.pop(), width=Inches(1.25))
        print("翻译完成")
        if not os.path.exists("tempfile"):
            os.makedirs("tempfile")
        document.save("tempfile/" + title_ch + ".docx")


    def output_html(self):
        fout = open('output.html', 'w',encoding='utf-8')
        fout.write("<html>")
        fout.write("<head>")
        fout.write(
            '<meta http-equiv="Content-Type" content="text/html; charset=utf-8" /> ')
        fout.write("</head>")
        fout.write("<body>")
        fout.write("<div>")

        for data in self.datas:
            for text in data['text']:
                fout.write("<div>")
                fout.write("<p>%s</p>" % text.encode('utf-8').decode('utf-8'))
                fout.write("</div>")
            fout.write("<br />")
        fout.write("</div>")
        fout.write("</body>")
        fout.write("</html>")
        fout.close()

    # def output_mysql(self):
        # db_config = {
        #     'host': '127.0.0.1',
        #     'port': 3306,
        #     'user': 'vini',
        #     'password': 'vini0926',
        #     'db': 'crawl_qidian',
        #     'charset': 'utf8',
        #     'cursorclass': pymysql.cursors.DictCursor,
        # }
        # connection = pymysql.connect(**db_config)
        # print (connection)
        # try:
        #     with connection.cursor() as cursor:
        #         sql = "insert into `novel_intro`(`name`,`intro`,`author`,`tags`,`type`,`message`,`word_count`,`click_count`,`groom_count`) values(s%,s%,s%,s%,s%,s%,s%,s%,s%)"
        #         cursor.execute(sql,
        #                        data['name'].encode('utf-8'),
        #                        data['intro'].encode('utf-8'),
        #                        data['author'].encode('utf-8'),
        #                        data['tags'].encode('utf-8'),
        #                        data['type'].encode('utf-8'),
        #                        data['message'].encode('utf-8'),
        #                        data['word_count'].encode('utf-8'),
        #                        data['click_count'].encode('utf-8'),
        #                        data['groom_count'].encode('utf-8'))

        #         connection.commit();
        # except expression as identifier:
        #     pass
        # finally:
        #     connection.close()
        # connection.close()
